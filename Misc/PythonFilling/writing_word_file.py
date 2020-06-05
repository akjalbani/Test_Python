# reference: automate boring stuff with python programming
# pip install python-docx
# Tested on python IDLE

import docx
import os
d = docx.Document()
d.add_paragraph("Hello This is first paragraph")
d.add_paragraph("This is second paragraph")
d.add_paragraph("Hello This is third paragraph")
d.add_paragraph("This is fourth paragraph")
# let us add bold text 
p = d.paragraphs[0] # will be added to the first line
p.add_run('This is new run')
p.runs[1].bold = True
d.save("C:\\Users\\xyz\\Desktop\\wordfiles\\demo5.docx")
